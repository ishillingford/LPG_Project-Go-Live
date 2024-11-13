import os
import zipfile
import extract_msg
import pandas as pd
from docx import Document
import re
from datetime import datetime
import json
from dotenv import load_dotenv
import openai 
from msal import PublicClientApplication
from office365.sharepoint.client_context import ClientContext 
from io import BytesIO
import os
import requests


# Load environment variables
load_dotenv()
openai.api_key = os.getenv('OPENAI_API_KEY')

# Define parameters
TENANT_ID = os.getenv("TENANT")  # Replace with your tenant ID
SHAREPOINT_SITE_URL = os.getenv("SHAREPOINT_SITE_URL")
GOLIVE_PATH = os.getenv("INPUT_PATH") 
OUTPUT_PATH= os.getenv("OUTPUT_PATH") 
UPLOAD_FOLDER_PATH = f"{SHAREPOINT_SITE_URL}/_api/web/GetFolderByServerRelativeUrl('{OUTPUT_PATH}')/Files/add(url='{{filename}}',overwrite=true)"

# Initialize MSAL app for interactive authentication
msal_app = PublicClientApplication(
    client_id="your-client-id-for-public-apps",
    authority=f"https://login.microsoftonline.com/{TENANT_ID}"
)

def authenticate_and_get_token():
    # Request access token interactively
    result = msal_app.acquire_token_interactive(scopes=["https://graph.microsoft.com/.default"])
    if "access_token" in result:
        print("Authentication successful!")
        return result["access_token"]
    else:
        print("Authentication failed.")
        return None

def fetch_files_from_sharepoint(folder_path):
    # Authenticate and get token
    token = authenticate_and_get_token()
    if token:
        headers = {"Authorization": f"Bearer {token}", "Accept": "application/json"}
        response = requests.get(
            f"{SHAREPOINT_SITE_URL}/_api/web/GetFolderByServerRelativeUrl('{folder_path}')/Files",
            headers=headers
        )
        if response.status_code == 200:
            files = response.json().get('d', {}).get('results', [])
            print("Files retrieved successfully!")
            return files
        else:
            print(f"Failed to retrieve files: {response.status_code} {response.text}")
            return None

def download_file_from_sharepoint(file_name, folder_path):
    token = authenticate_and_get_token()
    if token:
        headers = {"Authorization": f"Bearer {token}", "Accept": "application/json"}
        file_url = f"{SHAREPOINT_SITE_URL}/_api/web/GetFolderByServerRelativeUrl('{folder_path}')/Files('{file_name}')/$value"
        response = requests.get(file_url, headers=headers)
        if response.status_code == 200:
            return BytesIO(response.content)
        else:
            print(f"Failed to download file: {response.status_code} {response.text}")
            return None

def load_processed_emails():
    processed_emails_file = download_file_from_sharepoint('processed_emails.json', OUTPUT_PATH)
    if processed_emails_file:
        processed_emails = json.load(processed_emails_file)
        return set(processed_emails.get("processed_emails", []))
    else:
        return set()


def extract_info_from_msg(file_path): 
    msg = extract_msg.Message(file_path)
    email_date = msg.date
    email_subject = msg.subject

    # Check if email was processed already
    if email_subject in processed_emails:
        return None

    info = {
        "Project Title": "Not Provided",
        "Client Name": "Not Provided",
        "Use Case": "Not Provided",
        "Completion Date": "Not Provided",
        "Project Objectives": "Not Provided",
        "Business Challenges": "Not Provided",
        "Our Approach": "Not Provided",
        "Value Created": "Not Provided",
        "Measures of Success": "Not Provided",
        "Industry": "Not Provided"
    }

    body = re.sub(r'<[^>]+>', '', body)
    body = re.sub(r'\s+', ' ', body).strip()

    prompts = {
        "Project Title": "Extract the project title:",
        "Client Name": "Extract the client name (not Lionpoint):",
        "Use Case": "Extract the specific use case or objective of the project:",
        "Completion Date": "Extract the completion date (Month and Year):",
        "Project Objectives": "Extract the main objectives of the project:",
        "Business Challenges": "Extract the key business challenges faced by the client:",
        "Our Approach": "Extract the approach taken during the project:",
        "Value Created": "Extract the value created or outcomes achieved from the project:",
        "Measures of Success": "Extract the measures of success for the project:",
        "Industry": "Extract the industry related to the project:"
    }

    for key, prompt in prompts.items():
        response = openai.completions.create(
            model="gpt-3.5-turbo-instruct",
            prompt=f"{prompt}\n\n{body}",
            max_tokens=150,
            n=1,
            stop=None,
            temperature=0.5
        )
        info[key] = response.choices[0].text.strip()
        
        if info["Completion Date"] == "Not Provided" or not re.search(r'\b(?:January|February|March|April|May|June|July|August|September|October|November|December) \d{4}\b', info["Completion Date"], re.IGNORECASE):
            if email_date is not None:
                info["Completion Date"] = email_date.strftime("%B %Y")
            else:
                info["Completion Date"] = "Not Provided"
            
    return info

def add_heading_and_text(doc, heading, text, style=None):
    doc.add_heading(heading, level=2)
    if style:
        doc.add_paragraph(text, style=style)
    else:
        doc.add_paragraph(text)

def create_summary_doc(all_data, output_dir, example_doc_path):
    # Load the example document to match the style
    example_doc = Document(example_doc_path)
    example_styles = {style.name: style for style in example_doc.styles}

    doc = Document()
    doc.add_heading('Project Completion Summaries', 0)

    for project in all_data:
        doc.add_heading(project['Project Title'], level=1)
        
        add_heading_and_text(doc, 'Client Name:', project['Client Name'], example_styles.get('Body Text', None))
        add_heading_and_text(doc, 'Use Case:', project['Use Case'], example_styles.get('Body Text', None))
        add_heading_and_text(doc, 'Industry:', project['Industry'], example_styles.get('Body Text', None))
        add_heading_and_text(doc, 'Completion Date:', project['Completion Date'], example_styles.get('Body Text', None))
        
        add_heading_and_text(doc, 'Project Objectives:', project['Project Objectives'])
        add_heading_and_text(doc, 'Business Challenges:', project['Business Challenges'])
        add_heading_and_text(doc, 'Our Approach:', project['Our Approach'])
        add_heading_and_text(doc, 'Value Created:', project['Value Created'])
        add_heading_and_text(doc, 'Measures of Success:', project['Measures of Success'])
        
        doc.add_page_break()
    
    file_name = 'Project_Completion_Summaries_2.docx'
    doc.save(os.path.join(output_dir, file_name)) 
    
def summarize_info(info):
    summary_prompts = {
        "Project Objectives": "Summarize the project objectives briefly:",
        "Business Challenges": "Summarize the business challenges faced briefly:",
        "Our Approach": "Summarize our approach briefly:",
        "Value Created": "Summarize the value created briefly:",
        "Measures of Success": "Summarize the measures of success briefly:"
    }

    summarized_info = info.copy()

    for key, prompt in summary_prompts.items():
        if key in info:
            response = openai.completions.create(
                model="gpt-3.5-turbo-instruct",
                prompt=f"{prompt}\n\n{info[key]}",
                max_tokens=150,
                n=1,
                stop=None,
                temperature=0.5
            )
            summarized_info[key] = response.choices[0].text.strip()

    return summarized_info

def create_tracking_file(email_subjects):
    tracking_data = {"processed_emails": list(email_subjects)}
    tracking_file_bytes = BytesIO(json.dumps(tracking_data).encode())
    tracking_file_bytes.seek(0)
    upload_file_to_sharepoint(tracking_file_bytes, 'processed_emails.json')

def create_summary_doc(all_data):
    doc = Document()
    doc.add_heading("Project Completion Summaries", 0)

    for project in all_data:
        doc.add_heading(project['Project Title'], level=1)
        doc.add_paragraph(f"Client Name: {project['Client Name']}")
        doc.add_paragraph(f"Use Case: {project['Use Case']}")
        doc.add_paragraph(f"Industry: {project['Industry']}")
        doc.add_paragraph(f"Completion Date: {project['Completion Date']}")
        doc.add_paragraph("Project Objectives:", style='BodyText')
        doc.add_paragraph(project['Project Objectives'])
        doc.add_paragraph("Business Challenges:")
        doc.add_paragraph(project['Business Challenges'])
        doc.add_paragraph("Our Approach:")
        doc.add_paragraph(project['Our Approach'])
        doc.add_paragraph("Value Created:")
        doc.add_paragraph(project['Value Created'])
        doc.add_paragraph("Measures of Success:")
        doc.add_paragraph(project['Measures of Success'])
        doc.add_page_break()

    doc_bytes = BytesIO()
    doc.save(doc_bytes)
    doc_bytes.seek(0)
    upload_file_to_sharepoint(doc_bytes, 'Project_Completion_Summaries.docx')

def create_summary_excel(summarized_data):
    df = pd.DataFrame(summarized_data)
    excel_bytes = BytesIO()
    with pd.ExcelWriter(excel_bytes, engine="xlsxwriter") as writer:
        df.to_excel(writer, index=False)
    excel_bytes.seek(0)
    upload_file_to_sharepoint(excel_bytes, 'Project_Summaries.xlsx')

def upload_file_to_sharepoint(file_bytes, file_name):
    token = authenticate_and_get_token()
    if token:
        headers = {
            "Authorization": f"Bearer {token}",
            "Accept": "application/json;odata=verbose",
            "Content-Type": "application/octet-stream"
        }
        
        upload_url = UPLOAD_FOLDER_PATH.format(filename=file_name)
        response = requests.post(upload_url, headers=headers, data=file_bytes.getvalue())
        
        if response.status_code in [200, 201]:
            print(f"File '{file_name}' uploaded successfully!")
        else:
            print(f"Failed to upload '{file_name}': {response.status_code} {response.text}")

def main():
    files = fetch_files_from_sharepoint(GOLIVE_PATH)
    processed_emails = load_processed_emails()
    all_data = []
    summarized_data = []
    email_subjects = processed_emails.copy() 

    for file_info in files:
        file_name = file_info['Name']
        file_bytes = download_file_from_sharepoint(file_name, GOLIVE_PATH)
        if file_bytes and file_name.endswith('.msg'):
            with open(file_name, 'wb') as temp_file:
                temp_file.write(file_bytes.getvalue())
            info = extract_info_from_msg(file_name, processed_emails)
            if info:
                all_data.append(info)
                summarized_data.append(summarize_info(info))
                email_subjects.add(info['Project Title'])

    if all_data:
        create_summary_doc(all_data)
        create_summary_excel(summarized_data)
        create_tracking_file(email_subjects)

if __name__ == "__main__":
    main()

